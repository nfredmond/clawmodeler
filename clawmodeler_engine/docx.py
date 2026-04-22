from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from .report import render_report


class DocxDependencyMissingError(RuntimeError):
    pass


def _require_docx_deps():
    try:
        from docx import Document
        from markdown_it import MarkdownIt
    except ModuleNotFoundError as error:
        raise DocxDependencyMissingError(
            "DOCX export requires python-docx and markdown-it-py. Install the docx extra: "
            "`pip install clawmodeler-engine[docx]`."
        ) from error
    return Document, MarkdownIt


def render_docx(
    manifest: dict[str, Any],
    report_type: str,
    reports_dir: Path,
    *,
    ai_narrative: dict[str, Any] | None = None,
) -> bytes:
    """Render a report-type to DOCX bytes.

    The Markdown report produced by :func:`render_report` is parsed with
    markdown-it-py and walked into a python-docx ``Document``. Headings,
    paragraphs, inline emphasis, bullet and numbered lists, tables, and
    horizontal rules are preserved. Images are embedded when the file is
    resolvable relative to ``reports_dir`` and fall back to their alt text
    otherwise.
    """

    Document, MarkdownIt = _require_docx_deps()
    markdown = render_report(manifest, report_type, ai_narrative=ai_narrative)
    md = MarkdownIt("commonmark", {"html": False}).enable("table")
    tokens = md.parse(markdown)

    document = Document()
    _walk(document, tokens, reports_dir)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _walk(document, tokens, reports_dir: Path) -> None:
    i = 0
    while i < len(tokens):
        token = tokens[i]
        kind = token.type

        if kind == "heading_open":
            level = int(token.tag[1])  # "h2" -> 2
            inline = tokens[i + 1]
            heading = document.add_heading(level=min(level, 9))
            _apply_inline(heading, inline, reports_dir, embed_images=False)
            i += 3
            continue

        if kind == "paragraph_open":
            inline = tokens[i + 1]
            paragraph = document.add_paragraph()
            _apply_inline(paragraph, inline, reports_dir, embed_images=True)
            i += 3
            continue

        if kind == "bullet_list_open":
            i = _consume_list(document, tokens, i, "List Bullet", reports_dir)
            continue

        if kind == "ordered_list_open":
            i = _consume_list(document, tokens, i, "List Number", reports_dir)
            continue

        if kind == "table_open":
            i = _consume_table(document, tokens, i, reports_dir)
            continue

        if kind == "hr":
            document.add_paragraph()
            i += 1
            continue

        if kind == "blockquote_open":
            i = _consume_blockquote(document, tokens, i, reports_dir)
            continue

        if kind in ("fence", "code_block"):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(token.content.rstrip("\n"))
            run.font.name = "Consolas"
            i += 1
            continue

        i += 1


def _consume_list(document, tokens, start: int, style: str, reports_dir: Path) -> int:
    open_kind = tokens[start].type
    close_kind = open_kind.replace("_open", "_close")
    depth = 1
    i = start + 1
    while i < len(tokens) and depth > 0:
        token = tokens[i]
        if token.type == open_kind:
            depth += 1
        elif token.type == close_kind:
            depth -= 1
            if depth == 0:
                break
        elif token.type == "list_item_open" and depth == 1:
            i = _consume_list_item(document, tokens, i, style, reports_dir)
            continue
        i += 1
    return i + 1


def _consume_list_item(document, tokens, start: int, style: str, reports_dir: Path) -> int:
    depth = 1
    i = start + 1
    first_paragraph_written = False
    while i < len(tokens) and depth > 0:
        token = tokens[i]
        if token.type == "list_item_open":
            depth += 1
        elif token.type == "list_item_close":
            depth -= 1
            if depth == 0:
                break
        elif token.type == "paragraph_open" and depth == 1:
            inline = tokens[i + 1]
            if not first_paragraph_written:
                paragraph = document.add_paragraph(style=style)
                _apply_inline(paragraph, inline, reports_dir, embed_images=True)
                first_paragraph_written = True
            else:
                paragraph = document.add_paragraph()
                _apply_inline(paragraph, inline, reports_dir, embed_images=True)
            i += 3
            continue
        i += 1
    return i + 1


def _consume_table(document, tokens, start: int, reports_dir: Path) -> int:
    rows: list[list[Any]] = []
    current_row: list[Any] = []
    header_row_count = 0
    in_header = False
    i = start + 1
    depth = 1
    while i < len(tokens) and depth > 0:
        token = tokens[i]
        if token.type == "table_open":
            depth += 1
        elif token.type == "table_close":
            depth -= 1
            if depth == 0:
                break
        elif token.type == "thead_open":
            in_header = True
        elif token.type == "thead_close":
            in_header = False
        elif token.type == "tr_open":
            current_row = []
        elif token.type == "tr_close":
            rows.append(current_row)
            if in_header:
                header_row_count += 1
        elif token.type in ("th_open", "td_open"):
            inline = tokens[i + 1]
            current_row.append(inline)
            i += 3
            continue
        i += 1

    if rows:
        col_count = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=col_count)
        table.style = "Light Grid"
        for row_index, row in enumerate(rows):
            for col_index in range(col_count):
                cell = table.rows[row_index].cells[col_index]
                cell.text = ""  # python-docx seeds a default paragraph we overwrite
                paragraph = cell.paragraphs[0]
                if col_index < len(row):
                    _apply_inline(
                        paragraph,
                        row[col_index],
                        reports_dir,
                        embed_images=False,
                        bold=(row_index < header_row_count),
                    )
    return i + 1


def _consume_blockquote(document, tokens, start: int, reports_dir: Path) -> int:
    depth = 1
    i = start + 1
    while i < len(tokens) and depth > 0:
        token = tokens[i]
        if token.type == "blockquote_open":
            depth += 1
        elif token.type == "blockquote_close":
            depth -= 1
            if depth == 0:
                break
        elif token.type == "paragraph_open" and depth == 1:
            inline = tokens[i + 1]
            paragraph = document.add_paragraph(style="Intense Quote")
            _apply_inline(paragraph, inline, reports_dir, embed_images=True)
            i += 3
            continue
        i += 1
    return i + 1


def _apply_inline(
    paragraph,
    inline,
    reports_dir: Path,
    *,
    embed_images: bool,
    bold: bool = False,
) -> None:
    state = {"bold": bold, "italic": False, "link": None}
    for child in inline.children or []:
        kind = child.type
        if kind == "text":
            _add_run(paragraph, child.content, state)
        elif kind == "code_inline":
            _add_run(paragraph, child.content, state, force_code=True)
        elif kind == "softbreak":
            _add_run(paragraph, " ", state)
        elif kind == "hardbreak":
            paragraph.add_run().add_break()
        elif kind == "strong_open":
            state["bold"] = True
        elif kind == "strong_close":
            state["bold"] = bold
        elif kind == "em_open":
            state["italic"] = True
        elif kind == "em_close":
            state["italic"] = False
        elif kind == "link_open":
            state["link"] = child.attrs.get("href", "") if hasattr(child, "attrs") else ""
        elif kind == "link_close":
            if state["link"]:
                _add_run(paragraph, f" ({state['link']})", state)
            state["link"] = None
        elif kind == "image":
            alt = child.content or ""
            src = child.attrs.get("src", "") if hasattr(child, "attrs") else ""
            _add_image(paragraph, alt, src, reports_dir, embed_images)


def _add_run(paragraph, text: str, state: dict, *, force_code: bool = False) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    if state["bold"]:
        run.bold = True
    if state["italic"]:
        run.italic = True
    if force_code:
        run.font.name = "Consolas"


def _add_image(paragraph, alt: str, src: str, reports_dir: Path, embed: bool) -> None:
    if not embed or not src:
        paragraph.add_run(alt or src)
        return
    candidate = Path(src)
    if not candidate.is_absolute():
        candidate = (reports_dir / src).resolve()
    if not candidate.exists() or not candidate.is_file():
        paragraph.add_run(alt or src)
        return
    try:
        run = paragraph.add_run()
        run.add_picture(str(candidate))
    except Exception:
        paragraph.add_run(alt or src)
